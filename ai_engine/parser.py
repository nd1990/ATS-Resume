import os
import re
from pdfminer.high_level import extract_text
import docx

def clean_text(text):
    """
    Removes extra whitespace and cleans up the text.
    """
    if not text:
        return ""
    # Replace multiple newlines/spaces with single space
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def extract_text_from_pdf(pdf_path):
    try:
        text = extract_text(pdf_path)
        return clean_text(text)
    except Exception as e:
        print(f"Error reading PDF {pdf_path}: {e}")
        return ""

def extract_text_from_docx(docx_path):
    try:
        doc = docx.Document(docx_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return clean_text('\n'.join(full_text))
    except Exception as e:
        print(f"Error reading DOCX {docx_path}: {e}")
        return ""


def extract_text_from_image(image_path):
    try:
        from PIL import Image
        import pytesseract

        # Set tesseract path if not in PATH (common Windows paths)
        tesseract_cmd_paths = [
            r'C:\Program Files\Tesseract-OCR\tesseract.exe',
            r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
            r'D:\Resume ATS Checker\Resume ATS Checker\tesseract.exe',
            r'D:\jalak\Resume ATS Checker\ATS-Resume-Repo\tesseract.exe',
        ]

        # Check if tesseract is in PATH, if not try specific paths
        current_cmd = getattr(pytesseract.pytesseract, 'tesseract_cmd', '')
        if not current_cmd or not os.path.exists(current_cmd):
            for path in tesseract_cmd_paths:
                if os.path.exists(path):
                    pytesseract.pytesseract.tesseract_cmd = path
                    break
            else:
                # Tesseract not found — return empty string gracefully
                return ""

        # Set TESSDATA_PREFIX if not already set
        tessdata_paths = [
            r'D:\Resume ATS Checker\Resume ATS Checker\tessdata',
            r'D:\jalak\Resume ATS Checker\ATS-Resume-Repo\tessdata',
            r'C:\Program Files\Tesseract-OCR\tessdata',
            r'C:\Program Files (x86)\Tesseract-OCR\tessdata',
        ]
        if 'TESSDATA_PREFIX' not in os.environ:
            for path in tessdata_paths:
                if os.path.exists(os.path.join(path, 'eng.traineddata')):
                    os.environ['TESSDATA_PREFIX'] = path
                    break

        text = pytesseract.image_to_string(Image.open(image_path))
        return clean_text(text)
    except ImportError:
        return ""
    except OSError:
        # WinError 233: broken pipe / tesseract subprocess failed — return gracefully
        return ""
    except Exception:
        return ""

def extract_resume_text(file_path):
    """
    Detects file type and extracts text accordingly.
    """
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext == '.docx':
        return extract_text_from_docx(file_path)
    elif ext in ['.jpg', '.jpeg', '.png']:
        return extract_text_from_image(file_path)
    else:
        # TODO: Add .doc support if needed
        pass
    return ""


def extract_jd_text_from_upload(uploaded_file):
    """
    Extract text from an uploaded JD file (Django InMemoryUploadedFile / TemporaryUploadedFile).
    Supports: PDF, DOCX, JPG, JPEG, PNG.
    Returns extracted text string or empty string on failure.
    """
    import tempfile
    import shutil

    name = uploaded_file.name or ""
    ext = os.path.splitext(name)[1].lower()

    # Write to a temp file so existing helpers (path-based) can work
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
            uploaded_file.seek(0)
            shutil.copyfileobj(uploaded_file, tmp)
            tmp_path = tmp.name

        if ext == '.pdf':
            return extract_text_from_pdf(tmp_path)
        elif ext == '.docx':
            return extract_text_from_docx(tmp_path)
        elif ext in ['.jpg', '.jpeg', '.png']:
            return extract_text_from_image(tmp_path)
        else:
            return ""
    except Exception as e:
        print(f"Error extracting JD text from upload: {e}")
        return ""
    finally:
        try:
            os.remove(tmp_path)
        except Exception:
            pass

def extract_jd_text_raw(uploaded_file):
    """
    Like extract_jd_text_from_upload but preserves newlines — for job-title extraction.
    PDF text is NOT collapsed so heading structure is intact.
    """
    import tempfile, shutil
    name = uploaded_file.name or ""
    ext = os.path.splitext(name)[1].lower()
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
            uploaded_file.seek(0)
            shutil.copyfileobj(uploaded_file, tmp)
            tmp_path = tmp.name

        if ext == '.pdf':
            try:
                raw = extract_text(tmp_path)  # pdfminer — preserves newlines
                return raw or ""
            except Exception:
                return ""
        elif ext == '.docx':
            try:
                doc = docx.Document(tmp_path)
                return '\n'.join(p.text for p in doc.paragraphs)
            except Exception:
                return ""
        elif ext in ['.jpg', '.jpeg', '.png']:
            return extract_text_from_image(tmp_path)  # OCR already returns newlines
        return ""
    except Exception as e:
        print(f"extract_jd_text_raw error: {e}")
        return ""
    finally:
        try:
            os.remove(tmp_path)
        except Exception:
            pass
